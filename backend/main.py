import os
import json
import io
import re
from collections import deque
from typing import Any
from datetime import datetime
import requests
from google import genai
from google.genai import types
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
import docx
from pypdf import PdfReader
from PIL import Image
import pytesseract
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

load_dotenv()

# We configure GEMINI_API_KEY loaded from .env
api_key = os.getenv("GEMINI_API_KEY", "")
gemini_model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
ai_provider = os.getenv("AI_PROVIDER", "auto").strip().lower()
ollama_base_url = os.getenv("OLLAMA_BASE_URL", "http://127.0.0.1:11434").rstrip("/")
ollama_model = os.getenv("OLLAMA_MODEL", "qwen2.5:7b")
client = None
if not api_key:
    print("Warning: GEMINI_API_KEY is not set.")
else:
    client = genai.Client(api_key=api_key)

app = FastAPI(title="DocuSense Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

class EntityInfo(BaseModel):
    fact: str = Field(description="The extracted entity (Date, Name, Dollar Amount, etc.)")
    so_what: str = Field(description="Why does this fact matter? Consequence or analysis.")

class AnalyzeResponse(BaseModel):
    entities: list[EntityInfo] = Field(description="List of extracted entities and their implications")
    next_steps: list[str] = Field(description="3-5 tasks. Each must start with an action/imperative verb (e.g., Schedule, Draft, Verify).")
    source: str = Field(description="Analysis source, one of 'gemini', 'ollama', or 'fallback'")
    warning: str | None = Field(default=None, description="Optional warning if fallback mode was used.")
    document_type: str | None = Field(default=None, description="Detected document type such as contract, resume, or general")
    document_name: str | None = Field(default=None, description="Detected document name such as aadhaar, pan, rent agreement, or resume")
    document_category: str | None = Field(default=None, description="High-level category such as personal identification document, contract, resume, or general document")
    ats_score: int | None = Field(default=None, description="Estimated ATS score when the document type is resume")


class TranslateRequest(BaseModel):
    result: AnalyzeResponse
    target_language: str = Field(description="Target language code: en, hi, mr")


class ExportRequest(BaseModel):
    result: AnalyzeResponse
    document_name: str | None = Field(default=None, description="Original uploaded filename, optional")
    export_format: str = Field(description="Export format: docx or pdf")


class ComparisonFileSummary(BaseModel):
    filename: str
    document_type: str
    ats_score: int | None = None
    contract_score: int | None = None
    highlights: list[str] = Field(default_factory=list)


class ComparisonRow(BaseModel):
    attribute: str
    values_by_file: dict[str, str]
    difference: str
    assessment: str


class CompareResponse(BaseModel):
    document_type: str
    files: list[ComparisonFileSummary]
    comparison_rows: list[ComparisonRow]
    verdict: str
    warning: str | None = None


def analyze_response_to_dict(result: AnalyzeResponse) -> dict[str, Any]:
    if hasattr(result, "model_dump"):
        return result.model_dump()
    return result.dict()


def shorten_text(text: str, max_words: int) -> str:
    raw = (text or "").strip()
    if not raw:
        return raw
    words = raw.split()
    if len(words) <= max_words:
        return raw
    return " ".join(words[:max_words]).rstrip(".,;:") + "..."


def compact_analysis_payload(data: dict[str, Any]) -> dict[str, Any]:
    compacted = dict(data)
    compacted_entities: list[dict[str, str]] = []
    for entity in compacted.get("entities", []) or []:
        compacted_entities.append(
            {
                "fact": shorten_text(str(entity.get("fact", "")), 22),
                "so_what": shorten_text(str(entity.get("so_what", "")), 32),
            }
        )
    compacted["entities"] = compacted_entities
    compacted["next_steps"] = [
        shorten_text(str(step), 18) for step in (compacted.get("next_steps", []) or [])
    ]
    return compacted


def sanitize_filename(filename: str) -> str:
    cleaned = re.sub(r"[^\w\-. ]+", "", (filename or "").strip())
    cleaned = cleaned.replace(" ", "_")
    return cleaned or "analysis_report"


def build_report_lines(source_data: dict[str, Any], original_name: str | None = None) -> list[str]:
    lines: list[str] = []
    lines.append("DocuSense Analysis Export")
    lines.append(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if original_name:
        lines.append(f"Document: {original_name}")
    lines.append(f"Analysis source: {source_data.get('source') or 'unknown'}")
    warning = source_data.get("warning")
    if warning:
        lines.append(f"Warning: {warning}")
    lines.append("")
    lines.append("Extracted Entities:")
    entities = source_data.get("entities", []) or []
    if not entities:
        lines.append(" - No entities extracted.")
    for index, entity in enumerate(entities, start=1):
        lines.append(f"{index}. {entity.get('fact', '')}")
        lines.append(f"   Why it matters: {entity.get('so_what', '')}")
    lines.append("")
    lines.append("Actionable Next Steps:")
    steps = source_data.get("next_steps", []) or []
    if not steps:
        lines.append(" - No next steps generated.")
    for index, step in enumerate(steps, start=1):
        lines.append(f"{index}. {step}")
    return lines


def generate_docx_export(source_data: dict[str, Any], original_name: str | None = None) -> bytes:
    document = docx.Document()
    document.add_heading("DocuSense Analysis Export", level=1)
    document.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if original_name:
        document.add_paragraph(f"Document: {original_name}")
    document.add_paragraph(f"Analysis source: {source_data.get('source') or 'unknown'}")
    if source_data.get("warning"):
        document.add_paragraph(f"Warning: {source_data['warning']}")

    document.add_heading("Extracted Entities", level=2)
    entities = source_data.get("entities", []) or []
    if not entities:
        document.add_paragraph("No entities extracted.", style="List Bullet")
    for entity in entities:
        document.add_paragraph(str(entity.get("fact", "")), style="List Bullet")
        document.add_paragraph(f"Why it matters: {entity.get('so_what', '')}")

    document.add_heading("Actionable Next Steps", level=2)
    steps = source_data.get("next_steps", []) or []
    if not steps:
        document.add_paragraph("No next steps generated.", style="List Bullet")
    for step in steps:
        document.add_paragraph(str(step), style="List Number")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_pdf_export(source_data: dict[str, Any], original_name: str | None = None) -> bytes:
    packet = io.BytesIO()
    pdf_canvas = canvas.Canvas(packet, pagesize=A4)
    width, height = A4
    y = height - 40

    def write_line(text: str, font_name: str = "Helvetica", font_size: int = 10) -> None:
        nonlocal y
        if y < 40:
            pdf_canvas.showPage()
            y = height - 40
        pdf_canvas.setFont(font_name, font_size)
        pdf_canvas.drawString(40, y, text[:160])
        y -= 14

    for line in build_report_lines(source_data, original_name):
        if line == "DocuSense Analysis Export":
            write_line(line, font_name="Helvetica-Bold", font_size=14)
        elif line.endswith(":"):
            write_line(line, font_name="Helvetica-Bold", font_size=11)
        else:
            write_line(line)

    pdf_canvas.save()
    packet.seek(0)
    return packet.read()


def enrich_entities_with_contract_signals(entities: list[dict[str, str]], extracted_text: str) -> list[dict[str, str]]:
    text = extracted_text or ""
    lower_text = text.lower()
    enriched = list(entities or [])
    existing_facts = {str(item.get("fact", "")).strip().lower() for item in enriched if isinstance(item, dict)}

    def add_entity(fact: str, so_what: str) -> None:
        key = fact.strip().lower()
        if not key or key in existing_facts:
            return
        enriched.append({"fact": fact.strip(), "so_what": so_what.strip()})
        existing_facts.add(key)

    date_pattern = r"\b(?:\d{1,2}[/-]){1,2}\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b"
    money_pattern = r"(?:USD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?"

    # Agreement/signing date
    signed_match = re.search(r"(?:signed|executed|dated|agreement date)[^.:\n]*?(%s)" % date_pattern, text, flags=re.IGNORECASE)
    if signed_match:
        add_entity(
            f"Agreement signing/execution date: {signed_match.group(1)}",
            "This date establishes when legal obligations begin and can anchor payment and compliance timelines.",
        )

    # Agreement expiry/end date
    expiry_match = re.search(r"(?:expire|expires|expiry|end date|termination date)[^.:\n]*?(%s)" % date_pattern, text, flags=re.IGNORECASE)
    if expiry_match:
        add_entity(
            f"Agreement expiry date: {expiry_match.group(1)}",
            "Track this date for renewal, renegotiation, or exit decisions before deadlines pass.",
        )

    # Security deposit
    deposit_match = re.search(r"security deposit[^$\n]*?(%s)" % money_pattern, text, flags=re.IGNORECASE)
    if deposit_match:
        add_entity(
            f"Security deposit: {deposit_match.group(1)}",
            "Deposit terms impact cash flow and must be reconciled against refund/forfeit conditions.",
        )
    elif "security deposit" in lower_text:
        add_entity(
            "Security deposit mentioned",
            "Confirm the exact deposit amount and refund conditions to avoid disputes.",
        )

    def extract_rent_due_phrase(source_text: str) -> str | None:
        text_local = source_text or ""
        # Common forms: "on the 7th of every month", "7th day of each month", "on or before 7th"
        month_day_patterns = [
            r"\b(?:on\s+or\s+before\s+)?(?:the\s+)?(\d{1,2}(?:st|nd|rd|th)?)\s*(?:day\s+)?of\s+(?:every|each)\s+month\b",
            r"\b(?:on\s+)?(?:the\s+)?(\d{1,2}(?:st|nd|rd|th)?)\s*(?:day\s+)?(?:of\s+)?(?:every|each)\s+calendar\s+month\b",
            r"\b(?:on\s+or\s+before\s+)(\d{1,2}(?:st|nd|rd|th)?)\b",
        ]
        for pattern in month_day_patterns:
            match = re.search(pattern, text_local, flags=re.IGNORECASE)
            if match:
                return f"{match.group(1)} of every month"

        clause_match = re.search(r"(?:rent[^.\n]{0,220}?(?:due|payable)[^.\n]{0,220})", text_local, flags=re.IGNORECASE)
        if clause_match:
            clause = clause_match.group(0).strip()
            day_match = re.search(r"\b(\d{1,2}(?:st|nd|rd|th)?)\b", clause, flags=re.IGNORECASE)
            if day_match and re.search(r"\bmonth\b", clause, flags=re.IGNORECASE):
                return f"{day_match.group(1)} of every month"
            return clause
        return None

    # Rent due cadence/day
    due_clause = extract_rent_due_phrase(text)
    if due_clause:
        # Remove noisy model-produced rent due clause rows and replace with normalized one.
        enriched = [
            item for item in enriched
            if not (isinstance(item, dict) and "rent due" in str(item.get("fact", "")).strip().lower())
        ]
        existing_facts = {str(item.get("fact", "")).strip().lower() for item in enriched if isinstance(item, dict)}
        add_entity(
            f"Rent due date: {due_clause}",
            "This defines monthly payment cadence and late-risk exposure.",
        )

    # Add all distinct currency values (up to 8) for broader coverage.
    seen_amounts: set[str] = set()
    for amount in re.findall(money_pattern, text):
        normalized = amount.replace(" ", "")
        if normalized.lower() in seen_amounts:
            continue
        seen_amounts.add(normalized.lower())
        add_entity(
            f"Monetary amount referenced: {amount}",
            "Validate what this amount represents (rent, deposit, penalty, fee) and map to obligations.",
        )
        if len(seen_amounts) >= 8:
            break

    # Add all distinct dates (up to 10) for timeline completeness.
    seen_dates: set[str] = set()
    for date_value in re.findall(date_pattern, text, flags=re.IGNORECASE):
        key = date_value.lower()
        if key in seen_dates:
            continue
        seen_dates.add(key)
        add_entity(
            f"Date referenced: {date_value}",
            "Place this on the agreement timeline and verify whether it triggers an action.",
        )
        if len(seen_dates) >= 10:
            break

    return enriched


def extract_text_from_upload(file_bytes: bytes, mime_type: str | None) -> str:
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([para.text for para in doc.paragraphs if para.text])
    if mime_type == "text/plain":
        return file_bytes.decode("utf-8", errors="ignore")
    if mime_type == "application/pdf":
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text: list[str] = []
            for page in reader.pages:
                pages_text.append(page.extract_text() or "")
            return "\n".join([chunk for chunk in pages_text if chunk.strip()])
        except Exception as e:
            print(f"PDF extraction warning: {e}", flush=True)
            return ""
    if mime_type in {"image/png", "image/jpeg", "image/jpg", "image/webp"}:
        try:
            img = Image.open(io.BytesIO(file_bytes))
            return pytesseract.image_to_string(img).strip()
        except Exception as e:
            print(f"OCR extraction warning: {e}", flush=True)
            return ""
    return ""


def fallback_analysis(extracted_text: str, warning: str, requires_ai_parsing: bool = False, filename: str = "") -> AnalyzeResponse:
    text = extracted_text.strip()
    if not text and requires_ai_parsing:
        entities = [
            EntityInfo(
                fact="This file type needs AI parsing (PDF/Image) and AI is currently unavailable.",
                so_what="Local fallback cannot reliably extract entities from this file right now.",
            )
        ]
        next_steps = [
            "Retry after cooldown when AI provider is available.",
            "Upload a TXT or DOCX file for local fallback extraction.",
            "Verify your API quota and billing in Gemini dashboard.",
        ]
        data = {
            "entities": [item.model_dump() if hasattr(item, "model_dump") else item.dict() for item in entities],
            "next_steps": next_steps,
            "source": "fallback",
            "warning": warning,
        }
        data = enrich_document_metadata(data, text, filename)
        return AnalyzeResponse(**data)

    if not text:
        text = "No parseable text was found in the uploaded document."

    date_matches = re.findall(r"\b(?:\d{1,2}[/-]){1,2}\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b", text, flags=re.IGNORECASE)
    currency_matches = re.findall(r"(?:USD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text)

    entities: list[EntityInfo] = []
    if date_matches:
        entities.append(
            EntityInfo(
                fact=f"Dates found: {', '.join(date_matches[:10])}",
                so_what="These dates may represent deadlines, expirations, or milestones that require scheduling.",
            )
        )
    if currency_matches:
        entities.append(
            EntityInfo(
                fact=f"Amounts found: {', '.join(currency_matches[:10])}",
                so_what="Financial figures should be validated and tied to approvals or budget tracking.",
            )
        )

    due_value = None
    due_match = re.search(
        r"\b(?:on\s+or\s+before\s+)?(?:the\s+)?(\d{1,2}(?:st|nd|rd|th)?)\s*(?:day\s+)?of\s+(?:every|each)\s+month\b",
        text,
        flags=re.IGNORECASE,
    )
    if due_match:
        due_value = f"{due_match.group(1)} of every month"
    else:
        rent_clause_match = re.search(r"(?:rent[^.\n]{0,220}?(?:due|payable)[^.\n]{0,220})", text, flags=re.IGNORECASE)
        if rent_clause_match:
            clause = rent_clause_match.group(0).strip()
            day_match = re.search(r"\b(\d{1,2}(?:st|nd|rd|th)?)\b", clause, flags=re.IGNORECASE)
            if day_match and re.search(r"\bmonth\b", clause, flags=re.IGNORECASE):
                due_value = f"{day_match.group(1)} of every month"
            else:
                due_value = clause
    if due_value:
        entities.append(
            EntityInfo(
                fact=f"Rent due date: {due_value}",
                so_what="This clause determines monthly payment timing and potential late-payment risk.",
            )
        )

    deposit_match = re.search(r"security deposit[^$\n]*?((?:USD\\s*)?\\$\\s?\\d[\\d,]*(?:\\.\\d{1,2})?)", text, flags=re.IGNORECASE)
    if deposit_match:
        entities.append(
            EntityInfo(
                fact=f"Security deposit: {deposit_match.group(1)}",
                so_what="Deposit obligations should be verified against refund terms and damage clauses.",
            )
        )

    if not entities:
        entities.append(
            EntityInfo(
                fact=f"Document excerpt: {text[:120]}{'...' if len(text) > 120 else ''}",
                so_what="Review this excerpt and classify key entities manually if needed.",
            )
        )

    next_steps = [
        "Review the extracted entities and confirm their accuracy.",
        "Schedule any date-driven obligations in your project calendar.",
        "Verify monetary values against the source agreement or report.",
    ]
    data = {
        "entities": [item.model_dump() if hasattr(item, "model_dump") else item.dict() for item in entities],
        "next_steps": next_steps,
        "source": "fallback",
        "warning": warning,
    }
    data = enrich_document_metadata(data, text, filename)
    data = compact_analysis_payload(data)
    return AnalyzeResponse(**data)


def build_fallback_warning(error_text: str) -> str:
    retry_match = re.search(r"retry in\s+([0-9]+)", error_text, flags=re.IGNORECASE)
    retry_hint = f" Retry in ~{retry_match.group(1)}s." if retry_match else ""

    if "RESOURCE_EXHAUSTED" in error_text or "quota" in error_text.lower():
        return f"AI quota reached. Showing local fallback analysis.{retry_hint}"
    if "ollama" in error_text.lower():
        return "Ollama is unavailable. Showing local fallback analysis."
    if any(token in error_text.lower() for token in ["connection refused", "failed to establish a new connection", "max retries exceeded", "connection error"]):
        return "Could not connect to Ollama. Ensure Ollama is running on OLLAMA_BASE_URL and try again."
    if "api key" in error_text.lower():
        return "AI key missing or invalid. Showing local fallback analysis."
    return "AI provider temporarily unavailable. Showing local fallback analysis."


def maybe_extract_pdf_ocr(file_bytes: bytes) -> str:
    """
    Best-effort OCR for scanned PDFs. Requires optional dependencies (pdf2image + poppler).
    If unavailable, returns empty string and lets caller fall back.
    """
    try:
        from pdf2image import convert_from_bytes  # type: ignore
    except Exception:
        return ""

    try:
        images = convert_from_bytes(file_bytes, fmt="png", first_page=1, last_page=3)
        chunks: list[str] = []
        for img in images:
            text = pytesseract.image_to_string(img).strip()
            if text:
                chunks.append(text)
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def parse_json_from_text(raw_text: str) -> dict[str, Any]:
    text = raw_text.strip()
    if not text:
        raise ValueError("Empty model response text.")

    # Handle markdown fenced JSON commonly returned by local LLMs.
    fenced_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
    if fenced_match:
        return json.loads(fenced_match.group(1))

    # Try direct JSON first.
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Fallback: first JSON object substring.
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return json.loads(text[start:end + 1])

    raise ValueError("No JSON object found in model response.")


def normalize_language_code(language_code: str) -> str:
    normalized = (language_code or "").strip().lower()
    if normalized in {"en", "english"}:
        return "en"
    if normalized in {"hi", "hindi"}:
        return "hi"
    if normalized in {"mr", "marathi"}:
        return "mr"
    raise HTTPException(status_code=400, detail="Unsupported language. Use en, hi, or mr.")


def has_expected_script_text(result_payload: dict[str, Any], language_code: str) -> bool:
    if language_code == "en":
        return True
    pattern = r"[\u0900-\u097F]"
    text_parts: list[str] = []
    for entity in result_payload.get("entities", []) or []:
        text_parts.append(str(entity.get("fact", "")))
        text_parts.append(str(entity.get("so_what", "")))
    for step in result_payload.get("next_steps", []) or []:
        text_parts.append(str(step))
    combined = " ".join(text_parts)
    return bool(re.search(pattern, combined))


def translate_text_with_google(text: str, language_code: str) -> str:
    raw = (text or "").strip()
    if not raw:
        return text
    response = requests.post(
        "https://translate.googleapis.com/translate_a/single",
        params={
            "client": "gtx",
            "sl": "auto",
            "tl": language_code,
            "dt": "t",
        },
        data={"q": raw},
        timeout=20,
    )
    response.raise_for_status()
    payload = response.json()
    translated_chunks = payload[0] if isinstance(payload, list) and payload else []
    translated = "".join(
        chunk[0] for chunk in translated_chunks if isinstance(chunk, list) and chunk and chunk[0]
    )
    return translated or text


def translate_structured_payload_with_google(result: AnalyzeResponse, language_code: str) -> dict[str, Any]:
    source_data = analyze_response_to_dict(result)
    translated_entities: list[dict[str, str]] = []
    for entity in source_data.get("entities", []) or []:
        translated_entities.append(
            {
                "fact": translate_text_with_google(str(entity.get("fact", "")), language_code),
                "so_what": translate_text_with_google(str(entity.get("so_what", "")), language_code),
            }
        )

    translated_steps = [
        translate_text_with_google(str(step), language_code)
        for step in (source_data.get("next_steps", []) or [])
    ]

    translated_warning = source_data.get("warning")
    if translated_warning:
        translated_warning = translate_text_with_google(str(translated_warning), language_code)

    return {
        "entities": translated_entities,
        "next_steps": translated_steps,
        "source": source_data.get("source"),
        "warning": translated_warning,
        "document_type": source_data.get("document_type"),
        "document_name": source_data.get("document_name"),
        "document_category": source_data.get("document_category"),
        "ats_score": source_data.get("ats_score"),
    }


def detect_document_name(raw_text: str, filename: str = "") -> str:
    text = f"{raw_text or ''} {filename or ''}".lower()
    if any(token in text for token in ["aadhaar", "aadhar", "uidai", "unique identification authority of india"]):
        return "aadhaar"
    if any(token in text for token in ["permanent account number", "income tax department", "pan card", " pan ", " pan-", " pan_"]):
        return "pan"
    if any(token in text for token in ["driving licence", "driving license", "dl no", "dl number"]):
        return "driving licence"
    if any(token in text for token in ["passport", "republic of india", "passport no"]):
        return "passport"
    if any(token in text for token in ["rent agreement", "rental agreement", "lease deed", "leave and license"]):
        return "rent agreement"
    if any(token in text for token in ["offer letter", "employment offer"]):
        return "offer letter"
    if any(token in text for token in ["invoice", "tax invoice", "bill to"]):
        return "invoice"
    if "resume" in text or "curriculum vitae" in text:
        return "resume"
    if "contract" in text or "agreement" in text or "lease" in text:
        return "contract agreement"
    return "general document"


def resolve_document_category(doc_type: str, doc_name: str) -> str:
    name = (doc_name or "").lower()
    if name in {"aadhaar", "pan", "driving licence", "passport"}:
        return "personal identification document"
    if name in {"invoice"}:
        return "financial document"
    if name in {"offer letter"}:
        return "employment document"
    if doc_type == "contract" or "agreement" in name:
        return "contract"
    if doc_type == "resume":
        return "resume"
    return "general document"


def detect_document_type(raw_text: str, filename: str = "") -> str:
    text = f"{raw_text or ''} {filename or ''}".lower()
    if any(token in text for token in ["aadhaar", "aadhar", "uidai", "pan card", "passport", "driving licence", "driving license"]):
        return "identity_document"
    resume_signals = [
        "resume",
        "curriculum vitae",
        "professional summary",
        "work experience",
        "education",
        "skills",
        "projects",
        "certifications",
    ]
    if sum(1 for token in resume_signals if token in text) >= 2:
        return "resume"
    if any(token in text for token in ["agreement", "lease", "contract", "termination", "security deposit"]):
        return "contract"
    if any(token in text for token in ["invoice", "bill to", "amount due", "tax invoice"]):
        return "financial_document"
    return "general"


def estimate_ats_score(raw_text: str) -> int:
    text = (raw_text or "").lower()
    if not text.strip():
        return 55

    score = 42
    for token, points in {
        "skills": 10,
        "experience": 12,
        "education": 10,
        "projects": 8,
        "certifications": 6,
        "summary": 5,
    }.items():
        if token in text:
            score += points

    if re.search(r"[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}", text):
        score += 4
    if re.search(r"\b(20\d{2}|19\d{2})\b", text):
        score += 3
    if re.search(r"\b\d+%|\b\d+\s*(years|yrs)\b", text):
        score += 4
    if re.search(r"\bpython|java|javascript|react|node|sql|aws|docker\b", text):
        score += 6

    return max(45, min(98, score))


def enrich_document_metadata(data: dict[str, Any], extracted_text: str, filename: str = "") -> dict[str, Any]:
    doc_type = detect_document_type(extracted_text, filename)
    doc_name = detect_document_name(extracted_text, filename)
    category = resolve_document_category(doc_type, doc_name)

    data["document_type"] = doc_type
    data["document_name"] = doc_name
    data["document_category"] = category

    if doc_type == "resume":
        data["ats_score"] = estimate_ats_score(extracted_text)
    else:
        data["ats_score"] = None
    return data


def estimate_contract_score(raw_text: str) -> int:
    text = (raw_text or "").lower()
    if not text.strip():
        return 50

    score = 58
    positives = {
        "termination for convenience": 8,
        "cure period": 5,
        "notice period": 4,
        "indemnity": 4,
        "confidentiality": 4,
        "limitation of liability": 6,
        "dispute resolution": 3,
        "force majeure": 3,
    }
    risks = {
        "non-refundable": -6,
        "auto-renew": -5,
        "automatic renewal": -5,
        "sole discretion": -4,
        "unlimited liability": -9,
        "immediate termination": -6,
        "penalty": -4,
    }
    for token, points in positives.items():
        if token in text:
            score += points
    for token, points in risks.items():
        if token in text:
            score += points

    if re.search(r"\b\d{1,3}\s*(day|days)\s+notice\b", text):
        score += 3
    if re.search(r"(?:USD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text):
        score += 2

    return max(35, min(96, score))


def summarize_resume_metrics(raw_text: str) -> dict[str, str]:
    text = raw_text or ""
    lower = text.lower()
    email = re.search(r"[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}", lower)
    phone = re.search(r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?)\d{3,4}[-.\s]?\d{3,4}", text)
    skills_hits = re.findall(r"\bpython|java|javascript|react|node|sql|aws|docker|kubernetes|excel|power bi\b", lower)
    years_hits = re.findall(r"\b\d+\+?\s*(?:years|yrs)\b", lower)

    return {
        "Contact email": email.group(0) if email else "Not found",
        "Phone": phone.group(0) if phone else "Not found",
        "Skills keyword count": str(len(skills_hits)),
        "Years of experience mentions": str(len(years_hits)),
        "Education section": "Yes" if "education" in lower else "No",
        "Projects section": "Yes" if "projects" in lower else "No",
        "Certifications section": "Yes" if "certifications" in lower else "No",
    }


def summarize_contract_metrics(raw_text: str) -> dict[str, str]:
    text = raw_text or ""
    lower = text.lower()
    money_pattern = r"(?:USD|INR|Rs\.?|₹)?\s*\$?\s?\d[\d,]*(?:\.\d{1,2})?"
    date_pattern = (
        r"\b(?:\d{1,2}[./-]){1,2}\d{2,4}\b|"
        r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b"
    )
    amounts = re.findall(money_pattern, text)
    dates = re.findall(date_pattern, text, flags=re.IGNORECASE)
    clauses = [part.strip() for part in re.split(r"[\r\n]+|(?<=[.;])\s+", text) if part.strip()]

    def find_clause(patterns: list[str], max_len: int = 140) -> str | None:
        for clause in clauses:
            clause_lower = clause.lower()
            if any(re.search(pattern, clause_lower, flags=re.IGNORECASE) for pattern in patterns):
                snippet = re.sub(r"\s+", " ", clause).strip()
                return snippet[:max_len] + ("..." if len(snippet) > max_len else "")
        return None

    termination_clause = find_clause([r"\bterminat(?:e|ion|ed|es)\b", r"\bcancel(?:lation|led|s)?\b"])
    notice_clause = find_clause([r"\b\d{1,3}\s*(?:day|days|month|months)\s+notice\b", r"\bnotice\s+period\b"])
    payment_due_clause = find_clause(
        [
            r"\b(?:rent|license fee|licence fee|fee|payment|consideration)\b.*\b(?:due|payable|paid)\b",
            r"\bon\s+or\s+before\s+the?\s*\d{1,2}(?:st|nd|rd|th)?\b",
            r"\b(?:each|every)\s+month\b",
        ]
    )
    limitation_clause = find_clause(
        [
            r"\blimitation of liability\b",
            r"\bliability\b.*\b(?:limited|limit|capped|cap)\b",
            r"\bmaximum liability\b",
        ]
    )
    indemnity_clause = find_clause([r"\bindemn(?:ity|ification|ify|ifies|ified)\b", r"\bhold harmless\b"])
    confidentiality_clause = find_clause(
        [r"\bconfidential(?:ity)?\b", r"\bnon[-\s]?disclosure\b", r"\bconfidential information\b", r"\bnda\b"]
    )
    renewal_clause = find_clause(
        [r"\brenew(?:al|ed|s)?\b", r"\bauto[-\s]?renew(?:al)?\b", r"\bautomatic(?:ally)?\s+renew(?:ed|al)?\b", r"\bextend(?:ed|s|able)?\b"]
    )
    penalty_clause = find_clause([r"\bpenalt(?:y|ies)\b", r"\blate fee\b", r"\bliquidated damages?\b", r"\binterest\b.*\boverdue\b"])
    dispute_clause = find_clause([r"\bdispute resolution\b", r"\barbitration\b", r"\bjurisdiction\b", r"\bgoverning law\b", r"\bexclusive courts?\b"])

    return {
        "Monetary amounts found": ", ".join([value.strip() for value in amounts[:4]]) if amounts else "None found",
        "Date count": str(len(dates)),
        "Termination clause": termination_clause or "Not found",
        "Notice period": notice_clause or "Not found",
        "Payment due clause": payment_due_clause or "Not found",
        "Renewal language": renewal_clause or "Not found",
        "Limitation of liability": limitation_clause or "Not found",
        "Indemnity clause": indemnity_clause or "Not found",
        "Confidentiality clause": confidentiality_clause or "Not found",
        "Penalty/liquidated damages": penalty_clause or "Not found",
        "Dispute resolution/governing law": dispute_clause or "Not found",
    }


def compare_metric_maps(file_metrics: dict[str, dict[str, str]]) -> list[ComparisonRow]:
    all_keys: set[str] = set()
    for metric_map in file_metrics.values():
        all_keys.update(metric_map.keys())

    rows: list[ComparisonRow] = []
    for key in sorted(all_keys):
        values_by_file = {filename: metrics.get(key, "N/A") for filename, metrics in file_metrics.items()}
        distinct_values = {value.strip().lower() for value in values_by_file.values()}
        is_diff = len(distinct_values) > 1
        rows.append(
            ComparisonRow(
                attribute=key,
                values_by_file=values_by_file,
                difference="Different across files" if is_diff else "Mostly similar",
                assessment="Review this difference for decision impact." if is_diff else "Low decision impact.",
            )
        )
    return rows


def top_highlights_from_metrics(metrics: dict[str, str]) -> list[str]:
    items = []
    for key, value in metrics.items():
        if value and value not in {"No", "None found", "Not found"}:
            items.append(f"{key}: {value}")
        if len(items) >= 3:
            break
    return items or ["No strong highlights detected from extracted text."]


def translate_with_gemini(result: AnalyzeResponse, language_code: str) -> dict[str, Any]:
    if client is None:
        raise RuntimeError("Gemini client unavailable: missing GEMINI_API_KEY.")

    language_name = {"hi": "Hindi", "mr": "Marathi"}[language_code]
    system_instruction = (
        "You are a professional translator for legal and business content. "
        "Translate text naturally and accurately while preserving numbers, dates, amounts, and entity meaning."
    )
    user_prompt = (
        f"Translate this JSON content to {language_name}. "
        "Keep the JSON schema exactly identical with keys: entities, next_steps, source, warning, document_type, document_name, document_category, ats_score. "
        "Only translate human-readable sentence values. Keep source unchanged.\n\n"
        f"JSON:\n{json.dumps(analyze_response_to_dict(result), ensure_ascii=False)}"
    )
    response = client.models.generate_content(
        model=gemini_model,
        contents=[user_prompt],
        config=types.GenerateContentConfig(
            system_instruction=system_instruction,
            response_mime_type="application/json",
        ),
    )
    return json.loads(response.text)


def translate_with_ollama(result: AnalyzeResponse, language_code: str) -> dict[str, Any]:
    language_name = {"hi": "Hindi", "mr": "Marathi"}[language_code]
    source_payload = analyze_response_to_dict(result)
    prompt = (
        f"You are an expert translator. Translate every human-readable sentence value in this JSON into {language_name}. "
        "Do not keep English text unless it is a named entity or proper noun. "
        "Preserve numbers, dates, and currency symbols exactly. "
        "Return valid JSON only with this exact schema:\n"
        '{"entities":[{"fact":"string","so_what":"string"}],"next_steps":["string"],"source":"string","warning":"string|null","document_type":"string|null","document_name":"string|null","document_category":"string|null","ats_score":0}\n\n'
        f"INPUT JSON:\n{json.dumps(source_payload, ensure_ascii=False)}"
    )

    def has_devanagari_text(translated: dict[str, Any]) -> bool:
        text_parts: list[str] = []
        for entity in translated.get("entities", []) or []:
            text_parts.append(str(entity.get("fact", "")))
            text_parts.append(str(entity.get("so_what", "")))
        for step in translated.get("next_steps", []) or []:
            text_parts.append(str(step))
        combined = " ".join(text_parts)
        return bool(re.search(r"[\u0900-\u097F]", combined))

    last_error: Exception | None = None
    for _ in range(2):
        try:
            response = requests.post(
                f"{ollama_base_url}/api/generate",
                json={"model": ollama_model, "prompt": prompt, "stream": False, "format": "json"},
                timeout=120,
            )
            response.raise_for_status()
            payload = response.json()
            translated = parse_json_from_text((payload.get("response") or "").strip())
            if not has_devanagari_text(translated):
                raise ValueError(f"Translation output did not contain {language_name} script text.")
            return translated
        except Exception as error:
            last_error = error

    raise RuntimeError(f"Ollama translation failed: {last_error}")


def analyze_with_gemini(
    system_instruction: str,
    user_prompt: str,
    contents_payload: list[Any],
    filename: str = "",
    extracted_text: str = "",
) -> dict[str, Any]:
    if client is None:
        raise RuntimeError("Gemini client unavailable: missing GEMINI_API_KEY.")

    print(f"Submitting to Gemini model: {gemini_model}...", flush=True)
    response = client.models.generate_content(
        model=gemini_model,
        contents=contents_payload,
        config=types.GenerateContentConfig(
            system_instruction=system_instruction,
            response_mime_type="application/json"
        )
    )
    data = json.loads(response.text)
    data["entities"] = enrich_entities_with_contract_signals(data.get("entities", []), extracted_text or "")
    data["source"] = "gemini"
    data.setdefault("warning", None)
    data = enrich_document_metadata(
        data,
        extracted_text or "",
        filename,
    )
    data = compact_analysis_payload(data)
    return data


def analyze_with_ollama(system_instruction: str, user_prompt: str, extracted_text: str, filename: str = "") -> dict[str, Any]:
    if not extracted_text.strip():
        raise RuntimeError("Ollama requires extracted text input for this file type.")

    full_prompt = (
        f"{system_instruction.strip()}\n\n"
        "Return strictly valid JSON in this exact schema:\n"
        '{"entities":[{"fact":"string","so_what":"string"}],"next_steps":["string"],"warning":null,"document_type":"string|null","document_name":"string|null","document_category":"string|null","ats_score":0}\n\n'
        f"{user_prompt}\n\nDOCUMENT:\n{extracted_text}"
    )
    print(f"Submitting to Ollama model: {ollama_model}...", flush=True)
    response = requests.post(
        f"{ollama_base_url}/api/generate",
        json={"model": ollama_model, "prompt": full_prompt, "stream": False, "format": "json"},
        timeout=120,
    )
    response.raise_for_status()
    payload = response.json()
    raw_text = (payload.get("response") or "").strip()
    data = parse_json_from_text(raw_text)
    data["entities"] = enrich_entities_with_contract_signals(data.get("entities", []), extracted_text)
    data["source"] = "ollama"
    data.setdefault("warning", None)
    data = enrich_document_metadata(data, extracted_text, filename)
    data = compact_analysis_payload(data)
    return data

import time

# Global sliding-window limiter: max 20 requests per 60 seconds.
RATE_LIMIT_WINDOW_SECONDS = 60
MAX_REQUESTS_PER_WINDOW = 20
request_timestamps = deque()
gemini_translation_backoff_until = 0.0


def enforce_rate_limit() -> None:
    now = time.time()

    while request_timestamps and now - request_timestamps[0] >= RATE_LIMIT_WINDOW_SECONDS:
        request_timestamps.popleft()

    if len(request_timestamps) >= MAX_REQUESTS_PER_WINDOW:
        oldest_in_window = request_timestamps[0]
        retry_after = max(1, int(RATE_LIMIT_WINDOW_SECONDS - (now - oldest_in_window)))
        raise HTTPException(
            status_code=429,
            detail=f"Rate limit exceeded. Max {MAX_REQUESTS_PER_WINDOW} requests per minute. Retry in {retry_after}s.",
        )

    request_timestamps.append(now)

@app.post("/analyze", response_model=AnalyzeResponse)
def analyze_document(file: UploadFile = File(...)):
    enforce_rate_limit()

    print(f"--- Received Request ---", flush=True)

    system_instruction = """
    You are 'The Architect' from DocuSense AI.
    Your goal is to process unstructured documents (especially contracts/agreements) into comprehensive, structured, actionable insights.

    Insight Logic Protocol:
    1) Extraction: Scan the FULL document and extract all material entities, not just top highlights.
    2) Analysis: For each entity, explain the business/legal impact in 'so_what'.
    3) Action Synthesis: Create 5-8 concrete tasks, each starting with an imperative verb.

    Contract Coverage Requirements (when present):
    - Agreement/signing/execution date
    - Effective/start date
    - Expiry/termination date
    - Rent/fee/payment amounts
    - Security deposit amount and conditions
    - Payment due day/cadence (e.g., monthly on 7th)
    - Penalties, notice periods, renewal/termination obligations
    - Parties and property/asset references

    Return precise extracted facts from the text. Avoid generic summaries.
    """

    base_prompt = (
        "Analyze the full provided document using the Insight Logic Protocol. "
        "Do not limit to top 5 entities. Include all important contractual and financial points found.\n\n"
        "Return strictly JSON matching the response schema."
    )
    
    try:
        file_bytes = file.file.read()
        mime_type = file.content_type
        print(f"File loaded: {len(file_bytes)} bytes, Mime: {mime_type}")
        contents_payload = []
        extracted_text = extract_text_from_upload(file_bytes, mime_type)
        if (not extracted_text.strip()) and mime_type == "application/pdf":
            extracted_text = maybe_extract_pdf_ocr(file_bytes) or extracted_text
        detected_type = detect_document_type(extracted_text, file.filename or "")

        user_prompt = base_prompt
        if detected_type == "contract":
            user_prompt = (
                base_prompt
                + "\n\nContract-specific requirements: Ensure the entities list includes separate facts (with exact values/clauses when present) for: "
                + "agreement signing/execution date, effective/start date, expiry/end/termination date, rent/fees, security deposit, payment due day/cadence, "
                + "notice period, penalties/late fees, renewal/auto-renew terms, and party names. "
                + "Return at least 10 entities if the document contains these items."
            )

        # Text-based files are pre-extracted; other formats go in as native parts.
        if extracted_text:
            contents_payload.append(extracted_text)
            contents_payload.append(user_prompt)
        else:
            # For PDF, png, jpeg, etc. Gemini inherently processes them inline
            part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
            contents_payload.append(part)
            contents_payload.append(user_prompt)

        provider_mode = ai_provider if ai_provider in {"auto", "gemini", "ollama"} else "auto"
        # If Gemini isn't configured, prefer Ollama in auto mode.
        if provider_mode == "auto" and client is None:
            provider_mode = "ollama"

        if provider_mode == "gemini":
            data = analyze_with_gemini(
                system_instruction,
                user_prompt,
                contents_payload,
                file.filename or "",
                extracted_text=extracted_text,
            )
            return AnalyzeResponse(**data)
        if provider_mode == "ollama":
            data = analyze_with_ollama(system_instruction, user_prompt, extracted_text, file.filename or "")
            return AnalyzeResponse(**data)

        # auto mode: Gemini first, then Ollama (only when Gemini is configured)
        try:
            data = analyze_with_gemini(
                system_instruction,
                user_prompt,
                contents_payload,
                file.filename or "",
                extracted_text=extracted_text,
            )
            return AnalyzeResponse(**data)
        except Exception as gemini_error:
            print(f"Gemini failed in auto mode: {gemini_error}", flush=True)
            data = analyze_with_ollama(system_instruction, user_prompt, extracted_text, file.filename or "")
            data["warning"] = "Gemini unavailable, used Ollama fallback."
            return AnalyzeResponse(**data)
    except Exception as e:
        print(f"Error during analysis: {e}")
        error_text = str(e)
        print(f"Provider failure details: {error_text}", flush=True)
        warning = build_fallback_warning(error_text)
        return fallback_analysis(
            extracted_text if "extracted_text" in locals() else "",
            warning,
            requires_ai_parsing=not bool(extracted_text if "extracted_text" in locals() else ""),
            filename=file.filename if "file" in locals() else "",
        )

@app.get("/health")
def health_check():
    return {"status": "ok"}


@app.post("/compare", response_model=CompareResponse)
async def compare_documents(files: list[UploadFile] = File(...)):
    enforce_rate_limit()
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Upload at least 2 files to compare.")
    if len(files) > 5:
        raise HTTPException(status_code=400, detail="You can compare up to 5 files at a time.")

    parsed_docs: list[tuple[UploadFile, str, str]] = []
    doc_types: list[str] = []
    warnings: list[str] = []

    for file in files:
        file_bytes = await file.read()
        extracted_text = extract_text_from_upload(file_bytes, file.content_type)
        if not extracted_text.strip():
            warnings.append(f"Could not reliably parse text from {file.filename}.")
        doc_type = detect_document_type(extracted_text, file.filename or "")
        parsed_docs.append((file, extracted_text, doc_type))
        doc_types.append(doc_type)

    dominant_type = max(set(doc_types), key=doc_types.count)
    if any(doc_type != dominant_type for doc_type in doc_types):
        raise HTTPException(
            status_code=400,
            detail=f"All uploaded files must be the same document type. Found types: {', '.join(sorted(set(doc_types)))}.",
        )

    file_summaries: list[ComparisonFileSummary] = []
    file_metrics: dict[str, dict[str, str]] = {}

    for file, text, _ in parsed_docs:
        filename = file.filename or "unnamed-file"
        if dominant_type == "resume":
            ats = estimate_ats_score(text)
            metrics = summarize_resume_metrics(text)
            file_summaries.append(
                ComparisonFileSummary(
                    filename=filename,
                    document_type=dominant_type,
                    ats_score=ats,
                    highlights=top_highlights_from_metrics(metrics),
                )
            )
        elif dominant_type == "contract":
            contract_score = estimate_contract_score(text)
            metrics = summarize_contract_metrics(text)
            file_summaries.append(
                ComparisonFileSummary(
                    filename=filename,
                    document_type=dominant_type,
                    contract_score=contract_score,
                    highlights=top_highlights_from_metrics(metrics),
                )
            )
        else:
            metrics = {
                "Character length": str(len(text or "")),
                "Date mentions": str(len(re.findall(r"\b(?:\d{1,2}[/-]){1,2}\d{2,4}\b", text or ""))),
                "Currency mentions": str(len(re.findall(r"(?:USD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text or ""))),
            }
            file_summaries.append(
                ComparisonFileSummary(
                    filename=filename,
                    document_type=dominant_type,
                    highlights=top_highlights_from_metrics(metrics),
                )
            )
        file_metrics[filename] = metrics

    comparison_rows = compare_metric_maps(file_metrics)
    verdict = "Files are comparable, but no clear winner detected."
    if dominant_type == "resume":
        winner = max(file_summaries, key=lambda item: item.ats_score or 0)
        verdict = f"{winner.filename} appears stronger for ATS screening based on extracted structure and keyword coverage."
    elif dominant_type == "contract":
        winner = max(file_summaries, key=lambda item: item.contract_score or 0)
        verdict = f"{winner.filename} appears safer/more favorable based on clause coverage and lower-risk wording heuristics."

    return CompareResponse(
        document_type=dominant_type,
        files=file_summaries,
        comparison_rows=comparison_rows,
        verdict=verdict,
        warning=" ".join(warnings) if warnings else None,
    )


@app.post("/export")
def export_analysis_report(payload: ExportRequest):
    export_format = (payload.export_format or "").strip().lower()
    if export_format not in {"docx", "pdf"}:
        raise HTTPException(status_code=400, detail="Unsupported export format. Use 'docx' or 'pdf'.")

    source_data = analyze_response_to_dict(payload.result)
    base_name = sanitize_filename(payload.document_name or "analysis_report")

    if export_format == "docx":
        file_bytes = generate_docx_export(source_data, payload.document_name)
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base_name}_analysis.docx"},
        )

    file_bytes = generate_pdf_export(source_data, payload.document_name)
    return StreamingResponse(
        io.BytesIO(file_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={base_name}_analysis.pdf"},
    )


@app.post("/translate", response_model=AnalyzeResponse)
def translate_result(payload: TranslateRequest):
    global gemini_translation_backoff_until
    language_code = normalize_language_code(payload.target_language)
    if language_code == "en":
        return payload.result

    provider_mode = ai_provider if ai_provider in {"auto", "gemini", "ollama"} else "auto"
    original_warning = payload.result.warning
    try:
        if provider_mode == "gemini":
            translated = translate_with_gemini(payload.result, language_code)
        elif provider_mode == "ollama":
            translated = translate_with_ollama(payload.result, language_code)
        else:
            now = time.time()
            if now < gemini_translation_backoff_until:
                translated = translate_with_ollama(payload.result, language_code)
            else:
                try:
                    translated = translate_with_gemini(payload.result, language_code)
                except Exception as gemini_error:
                    print(f"Gemini translation failed in auto mode: {gemini_error}", flush=True)
                    retry_match = re.search(r"retry in\s+([0-9]+(?:\.[0-9]+)?)s", str(gemini_error), flags=re.IGNORECASE)
                    retry_seconds = float(retry_match.group(1)) if retry_match else 30.0
                    gemini_translation_backoff_until = time.time() + max(10.0, retry_seconds)
                    translated = translate_with_ollama(payload.result, language_code)

        translated["source"] = payload.result.source
        translated["warning"] = original_warning
        translated["document_type"] = payload.result.document_type
        translated["document_name"] = payload.result.document_name
        translated["document_category"] = payload.result.document_category
        translated["ats_score"] = payload.result.ats_score

        if not has_expected_script_text(translated, language_code):
            raise RuntimeError("Translated output did not contain expected Hindi/Marathi script text.")

        return AnalyzeResponse(**translated)
    except Exception as error:
        print(f"Translation failed: {error}", flush=True)
        try:
            translated = translate_structured_payload_with_google(payload.result, language_code)
            if not has_expected_script_text(translated, language_code):
                raise RuntimeError("Google translation output missing expected script text.")
            return AnalyzeResponse(**translated)
        except Exception as fallback_error:
            print(f"Google translation fallback failed: {fallback_error}", flush=True)
            raise HTTPException(
                status_code=503,
                detail="Translation is currently unavailable. Please retry in a few seconds.",
            )
